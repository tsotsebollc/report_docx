# -*- coding: utf-8 -*-

import logging
from odoo.tests import common

_logger = logging.getLogger(__name__)

try:
    from docx import Document
except ImportError:
    _logger.debug("Can not import python-docx`.")


class TestReportDocx(common.TransactionCase):
    @classmethod
    def setUpClass(cls):
        super().setUpClass()
        # Odoo 16-аас хойш DISABLED_MAIL_CONTEXT ашиглах шаардлагагүй болно
        DISABLED_MAIL_CONTEXT = {
            "tracking_disable": True,
            "mail_create_nolog": True,
            "mail_create_nosubscribe": True,
            "mail_notrack": True,
            "no_reset_password": True,
        }
        cls.env = cls.env(context=dict(cls.env.context, **DISABLED_MAIL_CONTEXT))
        report_object = cls.env["ir.actions.report"]
        cls.docx_report = cls.env["report.report_docx.abstract"].with_context(
            active_model="res.partner"
        )
        cls.report_name = "report_docx.partner_docx"
        cls.report = report_object._get_report_from_name(cls.report_name)
        cls.docs = cls.env["res.company"].search([], limit=1).partner_id

    def test_report(self):
        """ DOCX тайлангийн бүтцийг шалгах тест """
        report = self.report
        self.assertEqual(report.report_type, "docx")
        rep = report._render(self.docs.ids, {})
        doc = Document(BytesIO(rep[0]))
        
        # Эхний гарчгийг шалгах
        first_paragraph = doc.paragraphs[0].text
        self.assertEqual(first_paragraph, "Partner Report")

        # Харилцагчийн нэрийг шалгах
        partner_name = doc.paragraphs[1].text
        self.assertTrue(self.docs.name in partner_name)

    def test_id_retrieval(self):
        """ Тухайн тайланг дуудах ID-г зөв авч байгаа эсэхийг шалгана """

        # WebUI-аас wizard ашиглаж дуудах үед
        objs = self.docx_report._get_objs_for_report(
            False, {"context": {"active_ids": self.docs.ids}}
        )
        self.assertEqual(objs, self.docs)

        # Код дотроос дуудах үед
        objs = self.docx_report.with_context(
            active_ids=self.docs.ids
        )._get_objs_for_report(False, False)
        self.assertEqual(objs, self.docs)

        # WebUI-аас шууд дуудах үед
        objs = self.docx_report._get_objs_for_report(
            self.docs.ids, {"data": [self.report_name, self.report.report_type]}
        )
        self.assertEqual(objs, self.docs)

        # Тайлан үүсгэх үед
        objs = self.docx_report._get_objs_for_report(self.docs.ids, {})
        self.assertEqual(objs, self.docs)

    def test_currency_format(self):
        """ Мөнгөн тэмдэгтийн форматыг шалгах """
        mnt = self.env.ref("base.MNT")  # Монгол төгрөг
        self.assertEqual(
            self.docx_report._report_docx_currency_format(mnt), "#,##0 ₮"
        )
        usd = self.env.ref("base.USD")
        self.assertEqual(
            self.docx_report._report_docx_currency_format(usd), "$#,##0.00"
        )
        eur = self.env.ref("base.EUR")
        self.assertEqual(
            self.docx_report._report_docx_currency_format(eur), "#,##0.00 €"
        )
