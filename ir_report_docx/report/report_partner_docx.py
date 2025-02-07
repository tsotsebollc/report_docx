# Copyright 2017 Creu Blanca
# License AGPL-3.0 or later (https://www.gnu.org/licenses/agpl.html).

from odoo import models
from docx import Document

class PartnerDocx(models.AbstractModel):
    _name = "report.report_docx.partner_docx"
    _inherit = "report.report_docx.abstract"
    _description = "Partner DOCX Report"

    def generate_docx_report(self, doc, data, partners):
        """
        Generate DOCX report for partners
        """
        doc.add_heading("Partner Report", level=1)

        for partner in partners:
            doc.add_paragraph(f"Name: {partner.name}", style="Heading 2")
            doc.add_paragraph(f"Email: {partner.email or 'Not Available'}")
            doc.add_paragraph(f"Phone: {partner.phone or 'Not Available'}")
            doc.add_paragraph('-' * 30)  # Separator line
